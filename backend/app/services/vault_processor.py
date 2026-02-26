"""Vault document processing pipeline — text extraction + Claude analysis."""

import asyncio
import base64
import io
import json
import logging
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document as DocxDocument

from app.services.llm import get_user_settings
from app.services.vault_storage import VaultStorage

logger = logging.getLogger(__name__)

ANALYSIS_PROMPT = """You are a document analysis AI. Analyze the provided document text and return a JSON response with exactly this structure:
{
  "title": "A concise, descriptive title for this document",
  "doc_type": "one of: bank_statement, brokerage_statement, receipt, invoice, contract, lease, insurance, tax_document, medical, pay_stub, resume, research, notes, report, letter, legal, other",
  "summary": "2-3 sentence summary of the document's content and purpose",
  "entities": [
    {"name": "entity name", "type": "person|organization|account|amount|date|address|phone|email|other", "value": "extracted value if applicable"}
  ],
  "key_facts": ["Fact 1 extracted from the document", "Fact 2", ...]
}
Extract ALL meaningful entities — names, organizations, dollar amounts, dates, account numbers (last 4 digits only for security), addresses. For financial documents, always extract amounts, dates, and account references. Be thorough but concise."""


def _extract_text_pdf(file_path: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    pages: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)
    return "\n\n--- Page Break ---\n\n".join(pages)


def _extract_text_pdf_ocr(file_path: str) -> str:
    """Fallback: extract page images from PDF for vision API."""
    import fitz

    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def _extract_text_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_csv(file_path: str) -> str:
    """Generate a text summary of a CSV file."""
    import pandas as pd

    df = pd.read_csv(file_path, nrows=200)
    lines = [
        f"CSV with {len(df)} rows and {len(df.columns)} columns.",
        f"Columns: {', '.join(df.columns.tolist())}",
        "",
        "Sample rows:",
        df.head(10).to_string(index=False),
    ]
    desc = df.describe(include="all")
    if not desc.empty:
        lines.extend(["", "Statistics:", desc.to_string()])
    return "\n".join(lines)


def _extract_text_txt(file_path: str) -> str:
    """Read a plain text file."""
    return Path(file_path).read_text(encoding="utf-8", errors="replace")


def _file_to_base64(file_path: str) -> str:
    """Encode an image file to base64."""
    return base64.b64encode(Path(file_path).read_bytes()).decode("utf-8")


_NO_TEMPERATURE_MODELS = {"gpt-5", "gpt-5-mini", "gpt-5-nano", "gpt-5.2"}


def _get_litellm_params(settings: dict) -> dict[str, str]:
    """Return model and api_key for litellm based on provider setting."""
    provider = settings["provider"]
    if provider == "anthropic":
        return {
            "model": f"anthropic/{settings.get('anthropic_model', 'claude-sonnet-4-20250514')}",
            "api_key": settings["anthropic_api_key"],
        }
    if provider == "groq":
        model_name = settings.get("groq_model", "llama-3.3-70b-versatile")
        if not model_name.startswith("groq/"):
            model_name = f"groq/{model_name}"
        return {"model": model_name, "api_key": settings["groq_api_key"]}
    model_name = settings.get("openai_model", "gpt-4o")
    return {
        "model": f"openai/{model_name}",
        "api_key": settings["openai_api_key"],
    }


def _call_claude_analysis(text: str, is_image: bool = False, image_b64: str | None = None) -> dict[str, Any]:
    """Send content to the LLM via litellm for structured analysis."""
    from litellm import completion

    settings = get_user_settings()
    params = _get_litellm_params(settings)
    model_name = params["model"].split("/", 1)[-1]

    messages: list[dict[str, Any]] = []

    if is_image and image_b64:
        messages.append({
            "role": "user",
            "content": [
                {"type": "text", "text": ANALYSIS_PROMPT},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{image_b64}"},
                },
            ],
        })
    else:
        messages.append({
            "role": "user",
            "content": f"{ANALYSIS_PROMPT}\n\nDocument content:\n{text[:12000]}",
        })

    kwargs: dict[str, Any] = {
        "model": params["model"],
        "messages": messages,
        "api_key": params["api_key"],
        "max_completion_tokens": 2000,
    }
    if model_name not in _NO_TEMPERATURE_MODELS:
        kwargs["temperature"] = 0

    try:
        response = completion(**kwargs)
        raw = response.choices[0].message.content
        start = raw.find("{")
        end = raw.rfind("}") + 1
        if start >= 0 and end > start:
            return json.loads(raw[start:end])
        return {"title": "Untitled", "doc_type": "other", "summary": raw[:300], "entities": [], "key_facts": []}
    except Exception as e:
        logger.exception("LLM analysis failed")
        return {"title": "Untitled", "doc_type": "other", "summary": f"Analysis failed: {e}", "entities": [], "key_facts": []}


IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}


class VaultProcessor:
    """Processes uploaded documents through extraction and analysis."""

    def __init__(self, storage: VaultStorage, memory=None):
        self._storage = storage
        self._memory = memory

    async def process_document(self, doc_id: str) -> None:
        """Full processing pipeline for a document."""
        doc = await self._storage.get_document(doc_id)
        if not doc:
            logger.error("Document %s not found", doc_id)
            return

        await self._storage.update_document(doc_id, {"status": "processing"})

        try:
            file_path = doc["file_path"]
            file_type = doc["file_type"]

            text = await self._extract_text(file_path, file_type)

            is_image = file_type in IMAGE_EXTENSIONS
            image_b64 = None
            if is_image:
                image_b64 = await asyncio.to_thread(_file_to_base64, file_path)

            analysis = await asyncio.to_thread(
                _call_claude_analysis, text, is_image, image_b64
            )

            updates: dict[str, Any] = {
                "status": "completed",
                "raw_text": text,
                "title": analysis.get("title", "Untitled"),
                "doc_type": analysis.get("doc_type", "other"),
                "summary": analysis.get("summary", ""),
                "entities_json": json.dumps(analysis.get("entities", [])),
                "key_facts_json": json.dumps(analysis.get("key_facts", [])),
            }
            await self._storage.update_document(doc_id, updates)

            if self._memory:
                try:
                    await asyncio.to_thread(
                        self._memory.add_document,
                        doc_id=doc_id,
                        title=updates["title"],
                        summary=updates["summary"],
                        key_facts=analysis.get("key_facts", []),
                        entities=analysis.get("entities", []),
                        raw_text=text,
                        user_id=doc["user_email"],
                    )
                except Exception:
                    logger.warning("Mem0 storage failed for doc %s, continuing", doc_id)

            logger.info("Document %s processed successfully", doc_id)

        except Exception as e:
            logger.exception("Document processing failed for %s", doc_id)
            await self._storage.update_document(doc_id, {
                "status": "failed",
                "error_message": str(e),
            })

    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text based on file type."""
        extractors = {
            "pdf": self._extract_pdf,
            "docx": lambda p: asyncio.to_thread(_extract_text_docx, p),
            "doc": lambda p: asyncio.to_thread(_extract_text_docx, p),
            "csv": lambda p: asyncio.to_thread(_extract_text_csv, p),
            "txt": lambda p: asyncio.to_thread(_extract_text_txt, p),
        }

        if file_type in IMAGE_EXTENSIONS:
            return ""

        extractor = extractors.get(file_type)
        if extractor:
            return await extractor(file_path)

        return ""

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract PDF text, falling back to PyMuPDF if pdfplumber yields nothing."""
        text = await asyncio.to_thread(_extract_text_pdf, file_path)
        if not text.strip():
            text = await asyncio.to_thread(_extract_text_pdf_ocr, file_path)
        return text
